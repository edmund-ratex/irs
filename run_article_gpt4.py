import csv
import docx
import configparser
import docx
import time
import openai
from libs.logger import get_logger
import os
import re
import docx
import sys
import sendEmail
import subprocess

logger = get_logger(__name__)
regex = re.compile(r'[\\\\\/:*?"<>|#]')
openai.api_key = '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'  # la
# openai.api_key = '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' #ma
# articelCrawerDir = "."
docxStr = ".docx"
visitedArticlesFile = "./visited.csv"
article_url_str = "article_url:"
keywordsStr = "Keywords:"
promptTemplate = '''Read this article and generate 5 keywords at the 1st line starts with 'Keywords:' and an improved educational article outline (I., A., 1., etc) with a title, based only on content that fits an educational article. Ignore anything promoting other services, anything unclear because it references images that are unavailable, and anything about the original author or source:"TEXT HERE"'''
promptTemplateSection = '''We will go section by section. Follow your outline and generate the content for ####. Output should be in markdown format without including Roman numerals, letters, or numbers, with first level being H2. The punctuation marks of the output content are English punctuation instead of other languages' punctuation'''
promptKeywords = '''provide the top 5 keywords from the article based on relevance and estimated search volume. I don't need to know the estimated search volume, just give one list that excludes any text in parentheses, and use abbreviations where appropriate. Please provide the keywords and separate each keyword with a comma so I can easily copy and paste for internal linking purposes'''


class SectionsPromtAndContent:
    def __init__(self, prompt, section_outline):
        self.prompt = prompt
        self.section_outline = section_outline

    
def begin_crawling(articelCrawerDir):
    result = subprocess.run([f"{articelCrawerDir}/articelCrawer.exe"], capture_output=True, text=True)    
    # result = subprocess.run([f"{articelCrawerDir}/articelCrawer.exe"])
    print(result)


def start():
    try:
        articelCrawerDir = get_args(1)
        article_counts = int(get_args(2))
 		# begin_crawling(articelCrawerDir)        
          
        exclude_dirs = ["Crypto_in_a_Minute"]
        visitedArticles = read_file(
            visitedArticlesFile).strip('\n').split('\n')
        visitedArticles = [elem.strip().strip('"').strip("'")
                           for elem in visitedArticles]
        count = 0
        last_index = 0
        last_index_retry_time = 0
        
        directory = f'{articelCrawerDir}'
        dirName = os.path.basename(directory)
        current_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        current_time = regex.sub("_", current_time)
        csvPath = f"./academy_articles/article_{os.path.basename(directory)}_{current_time}.csv"
        row1 = ["title", "content", "outlines",
                "keywords", "raw_file_name"]
        append_to_csv(row1, csvPath)

        for subdir, dirs, files in os.walk(directory):
            # subdir: E:/WorkPlace/GO/src/articelCrawer/articles/phemex.com\Analysis
            if os.path.basename(subdir) in exclude_dirs:
                print('skip dir:' + os.path.basename(subdir))
                continue
            index = 0
            while index < len(files):
                try:
                    print(
                        f'---------------scan files in dir {subdir}: {index+1}/{len(files)}, generated count: {count+1}/{article_counts}-------------------')
                    file = files[index]
                    index += 1
                    if count >= article_counts:
                        break
                    filepath = subdir + os.sep + file
                    visitedFilePath = f'{dirName}/{file}'
                    if filepath.endswith(docxStr) and visitedFilePath not in visitedArticles:
                        doc = docx.Document(filepath)
                        article = '\n'.join(
                            [para.text for para in doc.paragraphs])

                        # article_url = ''
                        # contentSplits = article.split("\n")
                        # if len(contentSplits) > 1 and contentSplits[0].startswith(article_url_str):
                        #     article = article.replace(
                        #         contentSplits[0], '').strip()
                        #     article_url = contentSplits[0].replace(
                        #         article_url_str, '').strip()
                        # else:
                        #     logger.error(
                        #         f"caption content must start with 'caption_url:'----{file}")
                        #     continue

                        prompt = promptTemplate.replace(
                            'TEXT HERE', article)

                        arti = request_chatGPT(prompt, [])
                        # arti = "Keywords: Alchemy Pay, ACH token, crypto payments, fiat payments, payment gateway\n\nTitle: Understanding Alchemy Pay: Bridging the Gap Between Crypto and Fiat Payments\n\nI. Introduction\n    A. Overview of Alchemy Pay\n    B. Purpose of Alchemy Pay\n\nII. Alchemy Pay Features\n    A. Acceptance of fiat and crypto payments\n    B. Support for various business structures\n    C. Global operation and payment channels\n\nIII. ACH Token\n    A. Role of ACH token in Alchemy Pay ecosystem\n    B. Pledging system\n    C. DAO Governance\n\nIV. Conversion and Transaction Process\n    A. Instant conversion between crypto and fiat\n    B. Low fees and fast settlement\n    C. Integration with Chainlink's Price Feed Oracle\n\nV. Large Commercial Application\n    A. Support for complex account infrastructure\n    B. Support for complex transactions\n    C. Cross-border functionality\n\nVI. Alchemy Pay Leadership and Partnerships\n    A. Founders and key team members\n    B. Strategic and corporate partners\n\nVII. ACH Price History and Prediction\n    A. Price performance since debut\n    B. Factors influencing price\n\nVIII. Future of Alchemy Pay\n    A. Competition in the crypto payment gateway industry\n    B. Importance of partnerships for success\n\nIX. Conclusion\n    A. Alchemy Pay's potential for growth\n    B. Challenges and opportunities in the payment service industry"

                        splits = arti.split("\n")
                        title = ""
                        keywords = ""
                        for item in splits:
                            if item.startswith(keywordsStr):
                                keywords = item.replace(
                                    keywordsStr, "").strip()
                            if item.startswith("Title:") or item.startswith('title:'):
                                title = item
                                break

                        # content = arti.replace(title, "")
                        if title.startswith("Title:") or title.startswith('title:'):
                            title = title.lstrip("Title:")
                            title = title.lstrip("title:")
                            title = title.strip()
                        if title == '':
                            log_err(f"titile: {file}")
                            continue

                        if keywords == '':
                            log_err(f"keywords: {file}")
                            continue

                        gpt_msgs = [{"role": "user", "content": f"{prompt}"}, {
                            "role": "assistant", "content": f"{arti}"}]
                        # keywords = request_chatGPT(
                        #     promptKeywords, gpt_msgs)
                        sections = split_sectoin_by_roman_numeral_line(
                            arti)
                        final_result = ""
                        section_outlines = ""
                        prompt_and_outline_arr = gen_prompt_sections(
                            sections, 10)
                        for prompt_and_outline in prompt_and_outline_arr:
                            section_resp = request_chatGPT(
                                prompt_and_outline.prompt, gpt_msgs)
                            print(
                                f'---------------sub-progress: scan files in dir {subdir}: {index+1}/{len(files)}, generated count: {count+1}/{article_counts}-------------------')
                            final_result += f"\n\n{section_resp}"
                            section_outlines += prompt_and_outline.section_outline + \
                                '\n-------------------------------\n'
                        final_result = final_result.strip().replace('’', "'").replace(
                            '“', '"').replace('”', '"').replace('—', '-').replace('–', '-')
                        row = [title, final_result, section_outlines,
                                keywords, file]
                        append_to_csv(row, csvPath)
                        append_to_csv([visitedFilePath.strip('"')],
                                        visitedArticlesFile)
                        count += 1
                except Exception as e:
                    index -= 1
                    errMsg = f"err_msg:{e}, index:{index}, file:{file}"
                    print(errMsg)
                    log_err(errMsg)
                    time.sleep(1)
                    if last_index == index:
                        last_index_retry_time += 1
                    else:
                        last_index = index
                        last_index_retry_time = 0

                    if last_index_retry_time >= 3:
                        break
        sendEmail.send_file("academy CSV", csvPath)
        logger.info(f'send email academy csv succ: {csvPath} ')
    except Exception as e:
        errMsg = f":{e}"
        print(errMsg)
        log_err(errMsg)
        input()

        # outline_below_lines: 


def gen_prompt_sections(sections, outline_below_lines):
    roman_num_arrs = []
    sections_outline = ''
    line_num = 0
    prompt_arr = []
    max_section_num_at_once = 2

    for section in sections:
        roman_num = extract_roman_number(section)
        if roman_num == "":
            continue
        section_new = section.replace('\n\n', '\n').strip()
        cur_section_lines = section_new.count('\n') + 1
        if line_num + cur_section_lines <= outline_below_lines and len(roman_num_arrs) < max_section_num_at_once:
            roman_num_arrs.append(roman_num)
            sections_outline += section
            line_num += cur_section_lines
        else:
            # 
            if len(roman_num_arrs) > 0:
                temp_part_strs = [
                    f'Part {roman_num}' for roman_num in roman_num_arrs]
                elem = SectionsPromtAndContent(promptTemplateSection.replace(
                    '####', ', '.join(temp_part_strs)), sections_outline)
                prompt_arr.append(elem)
            roman_num_arrs = [roman_num]
            sections_outline = section
            line_num = cur_section_lines

    if len(roman_num_arrs) > 0:
        temp_part_strs = [f'Part {roman_num}' for roman_num in roman_num_arrs]
        elem = SectionsPromtAndContent(promptTemplateSection.replace(
            '####', ', '.join(temp_part_strs)), sections_outline)
        prompt_arr.append(elem)
    return prompt_arr


def pre_processing2(article):
    return split_article(article)


def extract_roman_number(text):
    regex = r"\b([IVX]+)\.\s"
    match = re.search(regex, text)
    if match:
        return match.group(1)
    else:
        print("No match found.")
        return ""


def split_sectoin_by_roman_numeral_line(article):
    # Define a regular expression to match lines starting with a Roman numeral and a period
    regex = r"^[IVX]+\."
    lines = article.strip().split("\n")

    result = []
    section = ""
    # Loop through the lines to find the matching line
    for line in lines:
        if re.match(regex, line):
            if section != "":
                result.append(section)
            section = line
        elif section != "":
            section += f'\n{line}'
    result.append(section)
    return result


def split_article(article_text):
    sections = []
    section = ''
    section_length = 0

    # 
    paragraphs = article_text.split('\n\n')

    for para in paragraphs:
        # 
        if is_heading(para):
            if section_length > 0:
                sections.append(section)
            section = para + '\n'
            section_length = len(section)
        else:
            para_length = len(para)
            # 
            if section_length + para_length > 3500:
                sections.append(section)
                section = para + '\n'
                section_length = len(section)
            else:
                section += para + '\n'
                section_length += para_length + 1

    # 
    if section_length > 0:
        sections.append(section)

    return sections


def is_heading(paragraph):
    return len(paragraph) < 40


def get_args():
    if len(sys.argv) == 2:
        return sys.argv[1]
    else:
        log_err("Please provide current dir name as command line argument.")


def read_config(dirName):
    config = configparser.ConfigParser()
    config.read(f'../{dirName}/promptConfig.ini')
    promptTemplate = config.get('prompt', 'promptTemplate')
    max_token_size = config.getint('param', 'max_token_size')
    return promptTemplate, max_token_size


def genFullPrompt(promptTemplate, keywordArr, removeDoubleQuotes):
    # 
    for keyword in keywordArr:
        if removeDoubleQuotes:
            promptTemplate = promptTemplate.replace('""', f'{keyword}', 1)
        else:
            promptTemplate = promptTemplate.replace('""', f'"{keyword}"', 1)
    return promptTemplate


def log_err(errMsg):
    openai.util.log_info(f'Log_Error: {errMsg}')


def append_to_csv(row, filePath):
    dir = os.path.dirname(filePath)
    if not os.path.exists(dir):
        os.makedirs(dir)

    # with open(filePath, "a", encoding='utf-8') as file:
    #     file.write(row)

    with open(filePath, 'a', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(row)


def read_visited_articles(dirName):
    content = read_file(f"../{dirName}/key_words.csv")
    content = content.strip('\n')
    keywordsList = content.split("\n")
    return keywordsList


def read_file(fileName):
    content = ''
    with open(fileName, "r", encoding='utf-8') as file:
        content = file.read()
    return content


def request_chatGPT(prompt, history_msgs):
    msgs = history_msgs.copy()
    msgs.append({"role": "user", "content": f"{prompt}"})
    completion = openai.ChatCompletion.create(
        # model="gpt-3.5-turbo",
        model="gpt-4",
        temperature=0.3,
        messages=msgs
    )

    content = completion.choices[0].message.content
    content = content.strip('\n')
    return content


def get_rubydex_glossary():
    doc = docx.Document('./RubyDex Glossary 2.docx')
    rgb_bold_contents = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold and run.font.color.rgb[0] == 51 and run.font.color.rgb[1] == 112 and run.font.color.rgb[2] == 255:
                rgb_bold_contents.append(run.text)

    return rgb_bold_contents


def get_args(index):
    if len(sys.argv) > index:
        return sys.argv[index]
    else:
        log_err("Please provide current dir name as command line argument.")


start()
input("")
